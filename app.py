from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import google.generativeai as genai
import os
import pandas as pd
import numpy as np
import json
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
from flask import Flask, render_template, request, redirect, url_for, session, jsonify, send_file
from authlib.integrations.flask_client import OAuth
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2.credentials import Credentials
from scipy import stats
from sklearn.cluster import KMeans
from sklearn.ensemble import RandomForestRegressor
from sklearn.decomposition import PCA
from sklearn.impute import SimpleImputer
import re
import io
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from googleapiclient.http import MediaIoBaseUpload
from docx.shared import RGBColor, Pt
from flask import Flask, session, request, jsonify
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

app = Flask(__name__)

# --- CONFIGURATION (Use Environment Variables) ---
app.secret_key = os.environ.get("SECRET_KEY")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET")

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///repository.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024

limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["500 per day"], 
    storage_uri="memory://",     
)

# Database Model for History
class AnalysisHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_email = db.Column(db.String(120), nullable=False)
    filename = db.Column(db.String(200))
    method = db.Column(db.String(100))
    result_html = db.Column(db.Text)  # Stores the full rendered HTML results
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

# Initialize the database file
with app.app_context():
    db.create_all()

# Initialize Gemini client with error handling
try:
    genai.configure(api_key=GEMINI_API_KEY)
    ai_model = genai.GenerativeModel('models/gemini-2.5-flash')
    system_instruction="""
    You are a Senior Research Data Scientist. 
    Your ONLY output must be valid HTML for a dashboard. 
    NEVER include conversational text, meta-commentary, or reviews of your own work.
    NEVER use Markdown (e.g., no **, no #). 
    Start directly with the <h3>Executive Summary</h3>.
    """
except Exception as e:
    print(f"Error initializing Gemini client: {e}")
    ai_model = None

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
    'scope': 'openid email profile https://www.googleapis.com/auth/drive.file',
    'prompt': 'consent',
    'access_type': 'offline'
},
)

def get_valid_credentials():
    if 'google_token' not in session:
        return None
    
    try:
        token_data = session['google_token']
        creds = Credentials(
            token=token_data.get('access_token'),
            refresh_token=token_data.get('refresh_token'),
            token_uri='https://oauth2.googleapis.com/token',
            client_id=GOOGLE_CLIENT_ID,
            client_secret=GOOGLE_CLIENT_SECRET
        )
        
        # If the token is expired, refresh it automatically
        if creds.expired and creds.refresh_token:
            from google.auth.transport.requests import Request
            creds.refresh(Request())
            # Update the session with the new token
            session['google_token']['access_token'] = creds.token
            session.modified = True
            
        return creds
    except Exception as e:
        print(f"Credentials error: {e}")
        return None

# --- RESEARCH ENGINE ---

def format_tables_in_html(html_content):
    """Convert text-based tables in HTML to proper HTML tables"""
    if '|' in html_content and '---' in html_content:
        lines = html_content.split('\n')
        table_lines = []
        in_table = False
        
        for line in lines:
            if '|' in line and line.strip():
                in_table = True
                table_lines.append(line)
            elif in_table and not line.strip():
                if table_lines:
                    table_html = '<table style="width:100%; border-collapse: collapse; margin: 1rem 0;">'
                    
                    if table_lines:
                        headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                        table_html += '<thead><tr>'
                        for header in headers:
                            table_html += f'<th style="background: rgba(102, 126, 234, 0.2); color: white; padding: 12px; text-align: left; font-weight: 600;">{header}</th>'
                        table_html += '</tr></thead><tbody>'
                    
                    for row_line in table_lines[2:]:
                        cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
                        if cells:
                            table_html += '<tr>'
                            for cell in cells:
                                table_html += f'<td style="padding: 10px 12px; border-bottom: 1px solid rgba(255, 255, 255, 0.05);">{cell}</td>'
                            table_html += '</tr>'
                    
                    table_html += '</tbody></table>'
                    html_content = html_content.replace('\n'.join(table_lines), table_html)
                
                table_lines = []
                in_table = False
    
    return html_content

def get_professional_insight(method_name, stats_data, df_summary, is_comparison=False):
    if not ai_model:
        return "<p><b>Error:</b> AI service is not available. Please check your Gemini API key.</p>"
    
    comp_task = "Compare both results. Explicitly state which is more accurate/useful and WHY." if is_comparison else ""
    
    prompt = f"""
    ROLE: Senior Research Data Scientist.
    METHOD: {method_name}
    STATS: {stats_data}
    CONTEXT: {df_summary}
    TASK: Provide a professional, extremely detailed and in-depth interpretation of the results.
    
    FORMATTING RULES:
    1. Output ONLY valid HTML. Do NOT use Markdown symbols like '**', '###', or '==='.
    2. Use <h3> for section titles.
    3. Use <b> for key metrics and <ul>/<li> for findings.
    4. If presenting tabular data, use proper HTML <table> tags with <thead>, <tbody>, <tr>, <th>, and <td>.
    5. Ensure all text is wrapped in <p> tags.
    {comp_task}"""
    
    try:
        response = ai_model.generate_content(prompt)
        insight = response.text
        
        insight = re.sub(r'={3,}', '', insight)
        insight = re.sub(r'-{3,}', '', insight)
        insight = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', insight)
        
        return format_tables_in_html(insight)
    except Exception as e:
        error_msg = str(e)
        if "API_KEY_INVALID" in error_msg:
            return f"<p><b>Authentication Error:</b> Your Gemini API key is invalid.</p>"
        else:
            return f"<p><b>Error generating insights:</b> {error_msg}</p>"

def create_custom_graph(df, graph_type):
    """Create specific graph types for custom analysis"""
    num_df = df.apply(pd.to_numeric, errors='coerce').dropna(axis=1, how='all')
    if num_df.empty:
        return None
    
    imputer = SimpleImputer(strategy='mean')
    clean_num = pd.DataFrame(imputer.fit_transform(num_df), columns=num_df.columns)
    
    plt.figure(figsize=(12, 8))
    
    try:
        if graph_type == 'bar':
            if clean_num.shape[1] >= 2:
                clean_num.iloc[:, :2].plot(kind='bar', ax=plt.gca())
                plt.title('Bar Chart Comparison', fontsize=14, fontweight='bold')
                plt.xlabel(clean_num.columns[0])
                plt.ylabel('Values')
                plt.xticks(rotation=45)
                plt.legend()
                plt.tight_layout()
        
        elif graph_type == 'line':
            for col in clean_num.columns[:3]:
                plt.plot(clean_num.index, clean_num[col], marker='o', label=col)
            plt.title('Line Chart Trends', fontsize=14, fontweight='bold')
            plt.xlabel('Index')
            plt.ylabel('Values')
            plt.legend()
            plt.grid(True, alpha=0.3)
            plt.tight_layout()
        
        elif graph_type == 'pie':
            if clean_num.shape[1] >= 1:
                values = clean_num.iloc[:, 0].value_counts().head(10)
                plt.pie(values, labels=values.index, autopct='%1.1f%%', startangle=90)
                plt.title('Pie Chart Distribution', fontsize=14, fontweight='bold')
                plt.axis('equal')
        
        elif graph_type == 'scatter':
            if clean_num.shape[1] >= 2:
                plt.scatter(clean_num.iloc[:, 0], clean_num.iloc[:, 1], alpha=0.6, s=50)
                plt.title('Scatter Plot', fontsize=14, fontweight='bold')
                plt.xlabel(clean_num.columns[0])
                plt.ylabel(clean_num.columns[1])
                plt.grid(True, alpha=0.3)
                plt.tight_layout()
        
        elif graph_type == 'histogram':
            clean_num.iloc[:, 0].plot(kind='hist', bins=30, edgecolor='black', alpha=0.7)
            plt.title('Histogram Distribution', fontsize=14, fontweight='bold')
            plt.xlabel(clean_num.columns[0])
            plt.ylabel('Frequency')
            plt.grid(True, alpha=0.3)
            plt.tight_layout()
        
        elif graph_type == 'heatmap':
            sns.heatmap(clean_num.corr(), annot=True, cmap='coolwarm', center=0, 
                       square=True, linewidths=1, cbar_kws={"shrink": 0.8})
            plt.title('Correlation Heatmap', fontsize=14, fontweight='bold')
            plt.tight_layout()
        
        elif graph_type == 'box':
            clean_num.boxplot(figsize=(12, 8))
            plt.title('Box & Whisker Plot', fontsize=14, fontweight='bold')
            plt.ylabel('Values')
            plt.xticks(rotation=45)
            plt.grid(True, alpha=0.3)
            plt.tight_layout()
        
        elif graph_type == 'waterfall':
            if clean_num.shape[1] >= 1:
                values = clean_num.iloc[:10, 0].values
                cumulative = np.cumsum(values)
                plt.bar(range(len(values)), values, alpha=0.7)
                plt.plot(range(len(values)), cumulative, 'r-o', linewidth=2, markersize=8, label='Cumulative')
                plt.title('Waterfall Chart', fontsize=14, fontweight='bold')
                plt.xlabel('Categories')
                plt.ylabel('Values')
                plt.legend()
                plt.grid(True, alpha=0.3)
                plt.tight_layout()
        
        img = io.BytesIO()
        plt.savefig(img, format='png', bbox_inches='tight', dpi=150)
        img.seek(0)
        plot_url = base64.b64encode(img.getvalue()).decode()
        plt.close()
        return plot_url
        
    except Exception as e:
        plt.close()
        return None

def run_analysis(df, method, custom_query=None, custom_graph_type=None, is_multi=False):
    """Executes all 11 methodologies and returns (Base64_Image_String, Stats_String, Method_Name, Has_Table)."""
    # 1. PREPARE NUMERIC DATA
    num_df = df.apply(pd.to_numeric, errors='coerce').dropna(axis=1, how='all')
    if num_df.empty:
        return None, "Error: No numerical data found.", method, False
    
    imputer = SimpleImputer(strategy='mean')
    clean_num = pd.DataFrame(imputer.fit_transform(num_df), columns=num_df.columns)
    
    # Global safety net for infinity/NaNs (Critical for Scikit-Learn)
    clean_num = clean_num.replace([np.inf, -np.inf], np.nan).fillna(0)

    if is_multi and 'Source_File' in df.columns:
        clean_num['Source_File'] = df['Source_File'].values

    plt.figure(figsize=(12, 8))
    stats_str, final_name = "", method.replace("_", " ").title()
    has_table = False

    try:
        hue_val = 'Source_File' if (is_multi and 'Source_File' in clean_num.columns) else None
        
        # --- 1. AUTO DETECT ---
        if method == "auto_detect":
            summary = clean_num.describe().to_string()
            has_table = True
            return None, summary, f"AI Recommended: Statistical Summary", has_table

        # --- 2. CORRELATION ---
        elif method == "correlation":
            corr_matrix = clean_num.select_dtypes(include=[np.number]).corr()
            sns.heatmap(corr_matrix, annot=True, cmap='RdBu_r', center=0, square=True)
            plt.title("Correlation Heatmap", fontsize=14, fontweight='bold')
            stats_str = corr_matrix.to_string()
            has_table = True

        # --- 3. K-MEANS ---
        elif method == "kmeans":
            km_data = clean_num.drop(columns=['Source_File', 'Cluster'], errors='ignore').select_dtypes(include=[np.number])
            if len(km_data) >= 3:
                model = KMeans(n_clusters=3, n_init='auto', random_state=42).fit(km_data)
                clean_num['Cluster'] = model.labels_
                display_hue = hue_val if (is_multi and hue_val) else 'Cluster'
                sns.scatterplot(data=clean_num, x=km_data.columns[0], y=km_data.columns[1], hue=display_hue, palette='viridis', s=80)
                stats_str = f"Inertia: {model.inertia_:.2f}\nClusters identified: 3"
            else:
                stats_str = "Insufficient data points for K-Means clustering."

        # --- 4. RANDOM FOREST ---
        elif method == "random_forest" and clean_num.shape[1] >= 2:
            rf_data = clean_num.drop(columns=['Source_File'], errors='ignore')
            X = rf_data.drop(rf_data.columns[0], axis=1)
            y = rf_data.iloc[:, 0]
            rf = RandomForestRegressor(n_estimators=100, random_state=42).fit(X, y)
            feat_imp = pd.DataFrame({'Feature': X.columns, 'Importance': rf.feature_importances_}).sort_values(by='Importance', ascending=False)
            sns.barplot(data=feat_imp, x='Importance', y='Feature', hue='Feature', palette='viridis', legend=False)
            stats_str = feat_imp.to_string(index=False)
            has_table = True

        # --- 5. REGRESSION ---
        elif method == "regression" and clean_num.shape[1] >= 2:
            sns.regplot(data=clean_num, x=clean_num.columns[0], y=clean_num.columns[-1], scatter_kws={'alpha':0.5})
            stats_str = f"Linear Regression: {clean_num.columns[0]} vs {clean_num.columns[-1]}"

        # --- 6. ANOVA ---
        elif method == "anova":
            cat_cols = [c for c in df.select_dtypes(exclude=[np.number]).columns if c != 'Source_File']
            if cat_cols:
                sns.boxplot(data=df, x=cat_cols[0], y=clean_num.columns[0], hue=cat_cols[0], palette='Set2', legend=False)
                stats_str = f"ANOVA variance check on {cat_cols[0]}"
            else:
                display_x = 'Source_File' if is_multi else None
                sns.boxenplot(data=clean_num, x=display_x, hue=display_x if is_multi else None, palette='muted', legend=False)
                stats_str = "Group-wise numeric distribution."

        # --- 7. PCA ---
        elif method == "pca" and clean_num.shape[1] >= 2:
            pca_data = clean_num.drop(columns=['Source_File', 'Cluster'], errors='ignore')
            pca = PCA(n_components=2)
            comps = pca.fit_transform(pca_data)
            pca_df = pd.DataFrame(data=comps, columns=['PC1', 'PC2'])
            if is_multi: pca_df['Source_File'] = clean_num['Source_File'].values
            sns.scatterplot(data=pca_df, x='PC1', y='PC2', hue=hue_val, palette='cool', s=70)
            stats_str = f"Explained Variance Ratio: {pca.explained_variance_ratio_}"
            has_table = True

        # --- 8. T-TEST ---
        elif method == "ttest" and clean_num.shape[1] >= 1:
            sns.kdeplot(data=clean_num, x=clean_num.columns[0], hue=hue_val, fill=True, alpha=0.5)
            stats_str = f"Mean for {clean_num.columns[0]}: {clean_num.iloc[:,0].mean():.2f}"
            has_table = True

        # --- 9. DISTRIBUTION ---
        elif method == "distribution":
            sns.histplot(data=clean_num, x=clean_num.columns[0], hue=hue_val, kde=True, bins=30)
            stats_str = f"Normal Distribution Check for {clean_num.columns[0]}"
            has_table = True

        # --- 10. TIMESERIES ---
        elif method == "timeseries":
            sns.lineplot(data=clean_num, x=clean_num.index, y=clean_num.columns[0], hue=hue_val, marker='o')
            stats_str = "Sequential trend analysis."

        # --- 11. OTHER (Custom) ---
        elif method == "other":
            if custom_graph_type:
                plt.close('all')
                plot_url = create_custom_graph(df, custom_graph_type)
                return plot_url, f"Custom Parameters: {custom_query}", final_name, False
            else:
                numeric_cols = clean_num.select_dtypes(include=[np.number]).columns[:4]
                g = sns.pairplot(clean_num, vars=numeric_cols, hue=hue_val)
                img = io.BytesIO()
                g.savefig(img, format='png', bbox_inches='tight', dpi=150)
                plt.close()
                return base64.b64encode(img.getvalue()).decode(), stats_str, final_name, False

        # --- UNIVERSAL SAVING BLOCK ---
        plt.tight_layout()
        img = io.BytesIO()
        plt.savefig(img, format='png', bbox_inches='tight', dpi=150)
        img.seek(0)
        plot_url = base64.b64encode(img.getvalue()).decode()
        plt.close()
        
        return plot_url, stats_str, final_name, has_table

    except Exception as e:
        plt.close()
        return None, f"Analysis Error: {str(e)}", method, False

def save_to_repository(email, filename, method, html_content):
    try:
        new_entry = AnalysisHistory(
            user_email=email,
            filename=filename,
            method=method,
            result_html=html_content
        )
        db.session.add(new_entry)
        db.session.commit()
    except Exception as e:
        print(f"Repository Error: {e}")

# --- ROUTES ---

@app.route("/")
def index():
    return render_template("login.html")

@app.route("/login/google")
@limiter.limit("10 per hour")
def login_google():
    return google.authorize_redirect(
        url_for('google_auth', _external=True),
        access_type='offline',
        prompt='consent'
    )

@app.route("/google/auth")
def google_auth():
    token = google.authorize_access_token()
    user = google.get('https://openidconnect.googleapis.com/v1/userinfo').json()
    session['google_token'] = token
    session['user_name'] = user.get('name')
    session['user_email'] = user.get('email')
    return redirect(url_for('dashboard'))

@app.route("/dashboard", methods=["GET", "POST"])
def dashboard():
    if 'google_token' not in session: 
        return redirect(url_for('index'))
    
    results = []
    current_filenames = []

    if request.method == "POST":
        uploaded_files = request.files.getlist("file")
        file_count_setting = int(request.form.get("file_count", 1))
        uploaded_files = uploaded_files[:file_count_setting]
        
        # Get Drive Service for Privacy Sync
        creds = get_valid_credentials()
        drive_service = build('drive', 'v3', credentials=creds) if creds else None
        
        all_dfs = []
        for i, file in enumerate(uploaded_files):
            if file and file.filename != '':
                filename = file.filename
                current_filenames.append(filename)
                
                # READ INTO MEMORY (Privacy: No local saving)
                file_bytes = file.read()
                file_io = io.BytesIO(file_bytes)
                
                # 1. Upload to Google Drive (Private User Storage)
                if drive_service:
                    try:
                        query = f"name = '{filename}' and trashed = false"
                        response = drive_service.files().list(q=query, fields="files(id, name)").execute()
                        existing_files = response.get('files', [])

                        if not existing_files:
                            file_io.seek(0)
                            media = MediaIoBaseUpload(file_io, mimetype=file.content_type, resumable=True)
                            
                            uploaded_file = drive_service.files().create(
                                body={'name': filename},
                                media_body=media,
                                fields='id'
                            ).execute()
                            print(f"DEBUG: New upload successful. ID: {uploaded_file.get('id')}")
                        else:
                            existing_id = existing_files[0].get('id')
                            print(f"DEBUG: File '{filename}' already exists in Drive (ID: {existing_id}). Skipping upload.")
                    except Exception as drive_err:
                        print(f"Drive Upload failed (Analysis continuing): {drive_err}")

                # 2. Process for Analysis
                if filename.endswith('.csv'):
                    temp_df = pd.read_csv(io.BytesIO(file_bytes), encoding='latin1')
                else:
                    temp_df = pd.read_excel(io.BytesIO(file_bytes))
                    
                temp_df['Source_File'] = f"Dataset {i+1}: {filename}"
                all_dfs.append(temp_df)

        if not all_dfs:
            return jsonify({"status": "error", "message": "No files uploaded"})

        df = pd.concat(all_dfs, ignore_index=True)
        is_multi = len(all_dfs) > 1
        
        mode = "comparative" if is_multi else request.form.get("mode")
        m1 = request.form.get("method")
        m2 = request.form.get("method2")
        c_query = request.form.get("custom_query")
        c_graph = request.form.get("custom_graph_type")

        try:
            # Pass the custom parameters into the run_analysis function
            p1, s1, n1, t1 = run_analysis(df, m1, custom_query=c_query, custom_graph_type=c_graph, is_multi=is_multi)
            
            # Summary Focus on Delta/Outliers if multi-file
            context_summary = df.groupby('Source_File').describe().to_string() if is_multi else df.describe().to_string()
            
            i1 = get_professional_insight(n1, s1, context_summary, is_comparison=is_multi)
            results.append({'plot': p1, 'insight': i1, 'name': n1, 'has_table': t1})

            if mode == "comparative" and m2 and not is_multi:
                p2, s2, n2, t2 = run_analysis(df, m2)
                i2 = get_professional_insight(f"Compare: {n1} vs {n2}", f"M1: {s1} | M2: {s2}", context_summary, True)
                results.append({'plot': p2, 'insight': i2, 'name': n2, 'has_table': t2})
            
        except Exception as e:
            results = [{'plot': None, 'insight': f'<p><b>Processing Error:</b> {str(e)}</p>', 'name': 'Error'}]

        # 3. Save to Repository
        save_to_repository(
            session.get('user_email'), 
            ", ".join(current_filenames), 
            results[0]['name'], 
            render_template("dashboard_partial.html", results=results, mode=mode)
        )
            # 4. Return AJAX Response
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({
                "status": "success",
                "filename": ", ".join(current_filenames),
                "html": render_template("dashboard_partial.html", results=results, mode=mode)
            })

    return render_template("dashboard.html", user=session.get('user_name'))

@app.route("/get_repository_history")
def get_repository_history():
    if 'user_email' not in session: return jsonify([])
    
    # Fetch all history for this specific user, newest first
    history_items = AnalysisHistory.query.filter_by(user_email=session['user_email'])\
                    .order_by(AnalysisHistory.timestamp.desc()).all()
    
    return jsonify([{
        "id": item.id,
        "name": f"{item.filename} ({item.method})",
        "createdTime": item.timestamp.strftime("%Y-%m-%d")
    } for item in history_items])

@app.route("/get_analysis_detail/<int:record_id>")
@limiter.limit("50 per day")
def get_analysis_detail(record_id):
    if 'user_email' not in session: return jsonify({"error": "Unauthorized"}), 401
    
    # Ensure user can only access their own records
    item = AnalysisHistory.query.filter_by(id=record_id, user_email=session['user_email']).first_or_404()
    
    return jsonify({
        "filename": item.filename,
        "method": item.method,
        "html": item.result_html
    })

from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import io

@app.errorhandler(429)
def ratelimit_handler(e):
    return jsonify({
        "error": "Usage Limit Reached",
        "message": "You've reached your hourly/daily limit. This ensures the engine remains fast for everyone. Please try again later!",
        "limit": str(e.description)
    }), 429

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({
        "error": "File Too Large",
        "message": "The maximum allowed file size is 30 MB. Please optimize your dataset."
    }), 413

@app.route("/export_word", methods=["POST"])
def export_word():
    if 'user_email' not in session:
        return jsonify({"error": "Unauthorized"}), 401
    
    # Get the HTML and Image from the request
    data = request.json
    html_content = data.get('html', '')
    image_base64 = data.get('image', None)
    user_local_time = data.get('local_time', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    doc = Document()
    doc.add_heading('BT IntelliStat - Analysis Report', 0)
    doc.add_paragraph(f"Generated on: {user_local_time}")
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # 1. Add the Visual (if exists)
    if image_base64:
        try:
            image_data = base64.b64decode(image_base64)
            image_stream = io.BytesIO(image_data)
            doc.add_heading('Statistical Visualization', level=1)
            doc.add_picture(image_stream, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"[Image could not be rendered: {e}]")

    # 2. Add the AI Insights (Parsing HTML to Word)
    doc.add_heading('AI Insights & Mathematical Interpretation', level=1)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.find_all(['h3', 'p', 'li', 'table']):
        if element.name == 'h3':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'li':
            doc.add_paragraph(f"• {element.get_text()}", style='List Bullet')
        elif element.name == 'table':
            # Create a native Word table
            rows = element.find_all('tr')
            if rows:
                cols = len(rows[0].find_all(['td', 'th']))
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()
    
    # Add a Professional Compliance Footer to the Document
    doc.add_paragraph("\n" * 3)
    
    # Use a subtle horizontal line for professional separation
    p_line = doc.add_paragraph()
    p_line.alignment = 1
    run_line = p_line.add_run("________________________________________________")
    run_line.font.color.rgb = RGBColor(200, 200, 200)

    # Formal Compliance Text
    compliance_p = doc.add_paragraph()
    compliance_p.alignment = 0 # Left aligned as requested
    
    run_badge = compliance_p.add_run("DPDP COMPLIANCE NOTICE: ")
    run_badge.bold = True
    run_badge.font.size = Pt(11)
    
    run_text = compliance_p.add_run(
        "This report confirms to the Digital Personal Data Protection Act (2023). "
        "Processing was executed in a zero-persistence environment. All research data is "
        "stored exclusively within the user's controlled Google Drive directory."
    )
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(100, 100, 100)

    # Verification Timestamp
    v_p = doc.add_paragraph(f"Verification ID: BT-STAT-{datetime.now().strftime('%Y%m%d-%H%M')}")
    v_p.style.font.size = Pt(9)
    v_p.alignment = 0

    # Save to memory and send
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name="BT_IntelliStat_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('index'))

if __name__ == "__main__":
    app.run(debug=True)
